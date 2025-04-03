import streamlit as st
from docx import Document
from docx.shared import Pt
from io import BytesIO


# === Display user email and logout/reset ===
if "user_email" not in st.session_state or not st.session_state.user_email:
    st.warning("Please return to the homepage and enter your email to start.")
    st.stop()

with st.sidebar:
    st.markdown(f"👤 **Logged in as:** `{st.session_state.user_email}`")
    if st.button("🔁 Log out / Reset"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.experimental_rerun()
st.title("📤 Export Agent")
st.markdown("Export your complete research paper in Journal of Business Research format.")

# Input paper title and author
title = st.text_input("Paper Title:")
author = st.text_input("Author Name (optional):")

# Sections in correct JBR order
sections = [
    "Abstract", "Introduction", "Literature Review",
    "Hypotheses/Framework", "Methodology", "Results",
    "Discussion", "Conclusion"
]

# Check that content exists
if "drafts" not in st.session_state:
    st.warning("No section drafts found. Please write or generate content in the Writing Agent first.")
    st.stop()

# Optional preview
if st.checkbox("📋 Preview all sections"):
    for section in sections:
        st.subheader(section)
        st.markdown(st.session_state.drafts.get(section, "_No content yet._"))

# Word export logic
def generate_docx(title, author, drafts):
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    if title:
        doc.add_heading(title, 0)
    if author:
        doc.add_paragraph(f"Author: {author}\n")

    for i, section in enumerate(sections):
        content = drafts.get(section, "")
        if not content.strip():
            continue
        if section == "Abstract":
            doc.add_heading("Abstract", level=1)
        else:
            doc.add_heading(f"{i}. {section}", level=1)

        for para in content.strip().split("\n\n"):
            doc.add_paragraph(para.strip())

    # Add References if present
    if "references" in st.session_state and st.session_state.references:
        doc.add_heading("References", level=1)
        for ref in st.session_state.references:
            doc.add_paragraph(ref)

    file = BytesIO()
    doc.save(file)
    file.seek(0)
    return file

# Export button
if st.button("📥 Download Full Paper (.docx)"):
    docx_file = generate_docx(title, author, st.session_state.drafts)
    st.download_button(
        label="Download Word Document",
        data=docx_file,
        file_name="jbr_research_draft.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
