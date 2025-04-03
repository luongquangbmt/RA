from utils.model_utils import call_llm
import streamlit as st


# === Display user email and logout/reset ===
if "user_email" not in st.session_state or not st.session_state.user_email:
    st.warning("Please return to the homepage and enter your email to start.")
    st.stop()

with st.sidebar:
    st.markdown(f"👤 **Logged in as:** `{st.session_state.user_email}`")
    if st.button("🔁 Log out / Reset"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.experimental_rerun()
st.title("✍️ Writing Agent")
st.markdown("Draft sections of your research paper based on the structured outline.")

# Ensure a structure plan exists in session state
if "structure_plan" not in st.session_state or not st.session_state.structure_plan:
    st.warning("Please generate and save a structure plan using the Structure Agent first.")
    st.stop()

# Display current structure plan
st.markdown("### 🏗️ Current Structure Plan")
st.text_area("Structure Plan:", value=st.session_state.structure_plan, height=150, disabled=True)

# Initialize drafts in session state if not present
if "drafts" not in st.session_state:
    st.session_state.drafts = {}

# Generate drafts for each section
sections = ["Introduction", "Literature Review", "Methodology", "Results", "Discussion", "Conclusion"]

for section in sections:
    if st.button(f"📝 Draft {section}"):
        with st.spinner(f"Writing the {section} section..."):
            prompt = f"""You are an academic research assistant tasked with drafting the '{section}' section of a research paper.

Based on the following structure plan:

{st.session_state.structure_plan}

Draft the '{section}' section with appropriate content and depth.

Use formal academic language.
"""
            try:
                response = call_llm(prompt).strip()
                st.session_state.drafts[section] = response
                st.markdown(f"### ✍️ Drafted {section} Section")
                st.text_area(f"{section} Section:", value=response, height=400, key=f"{section}_text")
            except Exception as e:
                st.error(f"Error: {e}")

# Save drafts to session
if st.session_state.get("drafts"):
    if st.button("✅ Save Drafts"):
        st.success("Drafts saved.")



import docx
from docx.shared import Inches
import base64

def save_docx(content, filename):
    doc = docx.Document()
    for section, text in content.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(text)
        if section in st.session_state.annotations:
            doc.add_paragraph("📝 Notes: " + st.session_state.annotations[section], style="Intense Quote")
    path = f"/tmp/{filename}.docx"
    doc.save(path)
    return path

if st.button("📤 Export Draft to DOCX"):
    path = save_docx(st.session_state.drafts, "research_draft")
    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="research_draft.docx">📥 Download your DOCX</a>'
        st.markdown(href, unsafe_allow_html=True)



# === In-Text Citation and Bibliography ===
if "citations" not in st.session_state:
    st.session_state.citations = {}

def generate_citekey(title):
    base = title.strip().split()[0].lower()
    suffix = ''.join(random.choices(string.ascii_lowercase, k=4))
    return f"{base}{suffix}"

st.markdown("### 🔖 Insert Citations")
for section in st.session_state.drafts:
    draft = st.session_state.drafts[section]
    cite_title = st.text_input(f"Title to cite in {section}:", key=f"cite_{section}")
    if st.button(f"➕ Insert citation into {section}", key=f"addcite_{section}"):
        cite_key = generate_citekey(cite_title)
        citation = f"({cite_key})"
        st.session_state.drafts[section] += " " + citation
        st.session_state.citations[cite_key] = cite_title
        st.success(f"Citation {citation} added to {section}.")

if st.session_state.citations:
    st.markdown("### 📚 Bibliography")
    for k, v in st.session_state.citations.items():
        st.markdown(f"- **{k}**: {v}")



# === AI Autocomplete ===
st.markdown("### 🤖 AI Autocomplete")
autocomplete_prompt = st.text_input("Start a sentence to complete:")
if st.button("✨ Autocomplete"):
    if autocomplete_prompt.strip():
        response = call_llm(f"Continue this sentence academically: '{autocomplete_prompt.strip()}'")
        st.text_area("📝 Suggestion", value=response.strip(), height=150)



# === APA Bibliography Formatter ===
def format_apa(entry):
    author = entry.get("author", "").strip()
    year = entry.get("year", "").strip()
    title = entry.get("title", "").strip()
    journal = entry.get("journal", "").strip()
    volume = entry.get("volume", "").strip()
    pages = entry.get("pages", "").strip()
    return f"{author} ({year}). {title}. *{journal}*, {volume}, {pages}."

if "apa_entries" not in st.session_state:
    st.session_state.apa_entries = []

st.markdown("### 📚 APA Bibliography Builder")
with st.expander("➕ Add new APA citation"):
    title = st.text_input("Title")
    author = st.text_input("Author(s)")
    year = st.text_input("Year")
    journal = st.text_input("Journal Name")
    volume = st.text_input("Volume")
    pages = st.text_input("Pages")

    if st.button("➕ Add APA Citation"):
        entry = {
            "title": title,
            "author": author,
            "year": year,
            "journal": journal,
            "volume": volume,
            "pages": pages
        }
        st.session_state.apa_entries.append(entry)
        st.success("APA citation added.")

if st.session_state.apa_entries:
    st.markdown("### ✅ Formatted APA Bibliography")
    for i, entry in enumerate(st.session_state.apa_entries, 1):
        st.markdown(f"{i}. {format_apa(entry)}")
