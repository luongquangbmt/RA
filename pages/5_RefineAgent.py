from utils.model_utils import call_llm
import streamlit as st


# === Display user email and logout/reset ===
if "user_email" not in st.session_state or not st.session_state.user_email:
    st.warning("Please return to the homepage and enter your email to start.")
    st.stop()

with st.sidebar:
    st.markdown(f"👤 **Logged in as:** `{st.session_state.user_email}`")
    if st.button("🔁 Log out / Reset"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.experimental_rerun()
st.title("🔍 Refine Agent")
st.markdown("Refine and enhance your drafted research paper sections.")

# Ensure drafts exist in session state
if "drafts" not in st.session_state or not st.session_state.drafts:
    st.warning("Please generate and save drafts using the Writing Agent first.")
    st.stop()

# Display current drafts
st.markdown("### 📝 Current Drafts")
for section, content in st.session_state.drafts.items():
    st.markdown(f"**{section}**")
    st.text_area(f"{section} Section:", value=content, height=150, key=f"{section}_draft", disabled=True)

# Select section to refine
section_to_refine = st.selectbox("Select a section to refine:", list(st.session_state.drafts.keys()))

# Refine the selected section
if st.button(f"🔧 Refine {section_to_refine} Section"):
    with st.spinner(f"Refining the {section_to_refine} section..."):
        prompt = f"""You are an academic research assistant tasked with refining the '{section_to_refine}' section of a research paper.

Current draft of the '{section_to_refine}' section:

{st.session_state.drafts[section_to_refine]}

Please improve the clarity, coherence, and academic rigor of this section. Ensure that the language is formal and that the content aligns with standard academic conventions.

Provide the refined version below:
"""
        try:
            response = call_llm(prompt).strip()
            st.session_state.drafts[section_to_refine] = response
            st.markdown(f"### ✨ Refined {section_to_refine} Section")
            st.text_area(f"{section_to_refine} Section:", value=response, height=400, key=f"{section_to_refine}_refined")
        except Exception as e:
            st.error(f"Error: {e}")

# Save refined drafts to session
if st.session_state.get("drafts"):
    if st.button("✅ Save Refined Drafts"):
        st.success("Refined drafts saved.")



import docx
from docx.shared import Inches
import base64

def save_docx(content, filename):
    doc = docx.Document()
    for section, text in content.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(text)
        if section in st.session_state.annotations:
            doc.add_paragraph("📝 Notes: " + st.session_state.annotations[section], style="Intense Quote")
    path = f"/tmp/{filename}.docx"
    doc.save(path)
    return path

if st.button("📤 Export Draft to DOCX"):
    path = save_docx(st.session_state.drafts, "research_draft")
    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="research_draft.docx">📥 Download your DOCX</a>'
        st.markdown(href, unsafe_allow_html=True)



# === In-Text Citation and Bibliography ===
if "citations" not in st.session_state:
    st.session_state.citations = {}

def generate_citekey(title):
    base = title.strip().split()[0].lower()
    suffix = ''.join(random.choices(string.ascii_lowercase, k=4))
    return f"{base}{suffix}"

st.markdown("### 🔖 Insert Citations")
for section in st.session_state.drafts:
    draft = st.session_state.drafts[section]
    cite_title = st.text_input(f"Title to cite in {section}:", key=f"cite_{section}")
    if st.button(f"➕ Insert citation into {section}", key=f"addcite_{section}"):
        cite_key = generate_citekey(cite_title)
        citation = f"({cite_key})"
        st.session_state.drafts[section] += " " + citation
        st.session_state.citations[cite_key] = cite_title
        st.success(f"Citation {citation} added to {section}.")

if st.session_state.citations:
    st.markdown("### 📚 Bibliography")
    for k, v in st.session_state.citations.items():
        st.markdown(f"- **{k}**: {v}")



# === AI Autocomplete ===
st.markdown("### 🤖 AI Autocomplete")
autocomplete_prompt = st.text_input("Start a sentence to complete:")
if st.button("✨ Autocomplete"):
    if autocomplete_prompt.strip():
        response = call_llm(f"Continue this sentence academically: '{autocomplete_prompt.strip()}'")
        st.text_area("📝 Suggestion", value=response.strip(), height=150)



# === APA Bibliography Formatter ===
def format_apa(entry):
    author = entry.get("author", "").strip()
    year = entry.get("year", "").strip()
    title = entry.get("title", "").strip()
    journal = entry.get("journal", "").strip()
    volume = entry.get("volume", "").strip()
    pages = entry.get("pages", "").strip()
    return f"{author} ({year}). {title}. *{journal}*, {volume}, {pages}."

if "apa_entries" not in st.session_state:
    st.session_state.apa_entries = []

st.markdown("### 📚 APA Bibliography Builder")
with st.expander("➕ Add new APA citation"):
    title = st.text_input("Title")
    author = st.text_input("Author(s)")
    year = st.text_input("Year")
    journal = st.text_input("Journal Name")
    volume = st.text_input("Volume")
    pages = st.text_input("Pages")

    if st.button("➕ Add APA Citation"):
        entry = {
            "title": title,
            "author": author,
            "year": year,
            "journal": journal,
            "volume": volume,
            "pages": pages
        }
        st.session_state.apa_entries.append(entry)
        st.success("APA citation added.")

if st.session_state.apa_entries:
    st.markdown("### ✅ Formatted APA Bibliography")
    for i, entry in enumerate(st.session_state.apa_entries, 1):
        st.markdown(f"{i}. {format_apa(entry)}")
