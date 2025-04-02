import streamlit as st
from huggingface_hub import InferenceClient
from docx import Document
from docx.shared import Pt
from io import BytesIO

# Streamlit page config
st.set_page_config(page_title="JBR Multi-Section Assistant", layout="wide")

# Hugging Face API setup
hf_token = st.secrets["HF_TOKEN"]
client = InferenceClient(model="mistralai/Mistral-7B-Instruct-v0.1", token=hf_token)

# Define sections
sections = [
    "Abstract", "Introduction", "Literature Review",
    "Hypotheses/Framework", "Methodology", "Results", "Discussion", "Conclusion"
]

# Title and author input
st.title("📄 JBR Research Paper Assistant – Multi-Section")
paper_title = st.text_input("Enter the title of your paper:")
author_name = st.text_input("Author name (optional):")

# Section selection
selected_section = st.selectbox("Choose a section to edit or generate:", sections)

# --- Citation Builder ---
st.markdown("---")
st.markdown("### 📚 Citation Builder")

# Initialize reference list
if "references" not in st.session_state:
    st.session_state.references = []

with st.expander("Add a New Reference"):
    author = st.text_input("Author(s):", key="author_input")
    title = st.text_input("Title of the article:", key="title_input")
    journal = st.text_input("Journal name:", key="journal_input")
    year = st.text_input("Publication year:", key="year_input")
    volume = st.text_input("Volume (optional):", key="volume_input")
    pages = st.text_input("Pages (optional):", key="pages_input")

    if st.button("➕ Add Reference"):
        if author and title and journal and year:
            citation = f"{author} ({year}). {title}. *{journal}*"
            if volume:
                citation += f", {volume}"
            if pages:
                citation += f", {pages}."
            else:
                citation += "."
            st.session_state.references.append(citation)
            st.success("Reference added.")
        else:
            st.error("Please fill in at least author, title, journal, and year.")

# Show current reference list
if st.session_state.references:
    st.markdown("#### 📖 Current References:")
    for i, ref in enumerate(st.session_state.references, start=1):
        st.markdown(f"{i}. {ref}")


# Initialize session state
if "paper" not in st.session_state:
    st.session_state.paper = {section: "" for section in sections}

# Input field for section content
st.markdown(f"### ✍️ {selected_section}")
user_input = st.text_area("Describe the topic, hypothesis, or method (for generation):", height=150)

# Generate section content using HF
if st.button("🤖 Generate Section"):
    with st.spinner("Generating..."):
        prompt = f"""You are a writing assistant specialized in academic business writing. Write the '{selected_section}' section of a Journal of Business Research article.

Topic/Notes: {user_input}

Write in formal academic tone with clarity and structure.
"""
        try:
            result = client.text_generation(prompt, max_new_tokens=500, temperature=0.7)
            st.session_state.paper[selected_section] = result.strip()
        except Exception as e:
            st.error(f"Error: {e}")

# Manual edit
st.session_state.paper[selected_section] = st.text_area(
    "Edit the content:", value=st.session_state.paper[selected_section], height=250
)

# Export DOCX function
def generate_full_docx(title, author, sections_dict):
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    if title:
        doc.add_heading(title, 0)
    if author:
        doc.add_paragraph(f"Author: {author}\n")

    for i, (section, content) in enumerate(sections_dict.items()):
        if section != "Abstract":
            doc.add_heading(f"{i}. {section}", level=1)
        else:
            doc.add_heading("Abstract", level=1)
        for para in content.strip().split("\n\n"):
            doc.add_paragraph(para.strip())

    # Add references at the end
    if st.session_state.references:
        doc.add_heading("References", level=1)
        for ref in st.session_state.references:
            doc.add_paragraph(ref)

    file = BytesIO()
    doc.save(file)
    file.seek(0)
    return file


# Export full paper
if st.button("📥 Download Full Paper (.docx)"):
    docx_file = generate_full_docx(paper_title, author_name, st.session_state.paper)
    st.download_button(
        label="Download Word Document",
        data=docx_file,
        file_name="jbr_draft.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
